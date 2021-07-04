
import requests
from bs4 import BeautifulSoup
import docx

filePath = 'C:\\Users\\dario\\Desktop\\epub.docx'
chapter = 1
chapterGoal = 1266
doc = docx.Document()
while chapter <= chapterGoal:
    link = 'https://bestlightnovel.com/novel_888154455/chapter_{}'.format(chapter)
    website = requests.get(link)
    text = BeautifulSoup(website.content, 'html.parser')
    chapterText = text.find('div', {'class' : 'vung_doc'})
    chapterContent = chapterText.find_all('p')
    doc.add_heading('Chapter {}\n'.format(chapter), 1)
    #f.write('Chapter {}\n'.format(chapter))
    for tag in chapterContent:
        doc.add_paragraph('   {}\n'.format(tag.get_text()))
    print(chapter)
    chapter += 1
doc.save(filePath)
